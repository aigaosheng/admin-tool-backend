import os
import sys
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()

import json
from typing import Dict, List, Optional, Union
from dataclasses import dataclass, asdict
import PyPDF2
import docx
from llms.llm_model import LlmModel
from audit_logger.audit_logger import logger_audit_handler as logger
from io import BytesIO
from docling.datamodel.base_models import DocumentStream
from pydantic import BaseModel
from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.document_converter import (
    DocumentConverter,
    PdfFormatOption,
    WordFormatOption,
)
from docling.pipeline.simple_pipeline import SimplePipeline
from docling.pipeline.standard_pdf_pipeline import StandardPdfPipeline

from util_tool.util import pil_image2base64
from doc_parser.docxMd import DocxMd
from doc_parser.paddleMd import PaddleMd

# Data Models
class TransactionModel(BaseModel):
    date_pickup: Optional[str] = None     
    id: Optional[str] = None 
    passenger: Optional[str] = None
    profile: Optional[str] = None
    location_pickup: Optional[str] = None
    destination: Optional[str] = None
    fare: Optional[float] = None
    platform_fee: Optional[float] = None
    driver_fee: Optional[float] = None
    total_paid: Optional[float] = None
    currency: Optional[str] = None

class TaxInvoiceModel(BaseModel):
    transactions: List[TransactionModel]    

class TaxInvoiceParser:
    def __init__(self, model_provider: str = "ollama", model_name: str = "gemma3"):
        """
        Initialize bank statement Parser with chosen model provider
        
        Args:
            model_provider: "ollama" or "gemini"
            model_name: Model name (for Ollama) or not used for Gemini
        """
        system_prompt2 = "You are a taxi invoice parser that extracts structured data from any quality data source. Missing field are set NULL. The taxi invoice is given in the below. Return as a JSON object."
        system_prompt = "Extract structured data from the image. Missing field are set NULL. The taxi invoice is given in the below. Return as a JSON object."

        self.model_provider = model_provider
        self.model_name = model_name
        self.llm = LlmModel(model_provider, model_name, system_prompt=system_prompt, response_data_model=TaxInvoiceModel)
        
        logger.info(f"LLM model: {model_provider}, {model_name}")
        
        self.inst_docmd = DocxMd()
        self.inst_paddle = PaddleMd()

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            if isinstance(file_path, bytes):
                buf = BytesIO(file_path)
                pdf_reader = PyPDF2.PdfReader(buf)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            else:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            logger.info(f"Success bank statement to text -> {file_path[:50]} -> {text[:50]}")
        except Exception as e:
            logger.warning(f"Error reading PDF: {e}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            if isinstance(file_path, bytes):
                buf = BytesIO(file_path)
                doc = docx.Document(buf)
            else:
                doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            logger.info(f"Success bank statement to text -> {file_path[:50]} -> {text[:50]}")
        except Exception as e:
            logger.warning(f"Error reading DOCX: {e}")
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            if isinstance(file_path, bytes):
                text = file_path.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            logger.info(f"Success bank statement to text -> {file_path[:50]} -> {text[:50]}")

            return text
        except Exception as e:
            logger.warning(f"Error reading TXT: {e}")
            return ""
    
    def extract_text_docling(self, file_path: str) -> str:
        """Extract text from resume file based on extension"""
        converter = DocumentConverter()  # all of the below is optional, has internal defaults.
        try:
            if isinstance(file_path, bytes):
                buf = BytesIO(file_path)
                source = DocumentStream(name="my_doc", stream=buf)
            else:
                source = file_path
            result = converter.convert(source)

            md = result.document.export_to_markdown()
            logger.info(f"Success bank statement to text -> {file_path}")
        except Exception as e:
            logger.warning(f"Error proess bank statement to text -> {file_path}")
            md = ""  
        return md

    def extract_text(self, file_path: str, name: str = "doctmp", doc_type = "image") -> str:
        """Extract text from resume file based on extension
        return:
        key-value dict. key = file name
        """
        if isinstance(file_path, bytes):
            if doc_type in ("image", "pdf"):
                logger.info(f"*** Paddlepaddle parser")
                result = self.inst_paddle.documentText(file_path, name, doc_type)
                # md = dict(map(lambda x: (x['page'], x['text']), result))
                md = result
            elif doc_type == 'doc':
                logger.info(f"*** DOC + Paddlepaddle parser")
                result = self.inst_docmd.documentText(file_path, name)
                md = result
            else:
                result = self.extract_text_docling(file_path)
                md = {"file_path": result }
        else:
            if Path(file_path).suffix.lower() in ['.jpg', '.jpeg', '.png', '.pdf']:
                result = self.inst_paddle.documentText(file_path, file_path, doc_type)
                # md = "\n\n".join(list(map(lambda x: f"## Document-{x[0]+1}\n" + x[1]["text"], enumerate(result))))
                # md = list(map(lambda x: x[1]["text"], enumerate(result)))
                md = dict(map(lambda x: (x['page'], x['text']), result))
            elif Path(file_path).suffix.lower() in ['.docx', 'doc']:
                result = self.inst_docmd.documentText(file_path, file_path)
                md = result
                # md = "\n\n".join(list(map(lambda x: x[1], result.items())))
                # md = list(map(lambda x: x[1], result.items()))
            else:
                result = self.extract_text_docling(file_path)
                md = {"file_path": result }

        return md

    async def adocumentParser(self, file_path: str) -> TaxInvoiceModel:
        """
        Parse resume and extract metadata
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            TaxInvoiceModel object with extracted information
        """
        # Extract text from resume
        resume_text = self.extract_text(file_path)
        # resume_text = self.extract_text_from_pdf(file_path)
        
        if not resume_text.strip():
            logger.warning(f"Could not extract text from resume file -> {file_path}")
        
        # Extract information
        try:
            result = await self.llm(resume_text)
            # rsp = self.flatten_meta(result)
            return result
        except Exception as e:
            logger.warning(f"Error during parsing: {e}")
            return None
    
    def documentParser(self, file_path: str, name: str, doc_type: str) -> TaxInvoiceModel:
        """
        Parse resume and extract metadata
        
        Args:
            file_path: Path to the resume file
            doc_type: "image", "pdf", "doc", other
            
        Returns:
            TaxInvoiceModel object with extracted information
        """
        # Extract text from resume
        resume_texts = self.extract_text(file_path, name, doc_type)
        # resume_text = self.extract_text_from_pdf(file_path)

        if isinstance(resume_texts, str):
            resume_texts = [{"name": name, "page": "page_text", "text": resume_texts, "image": None}]
        
        outputs = []
        for resume_text in resume_texts:
            if not resume_text["text"] or not resume_text["text"].strip():
                logger.warning(f"Could not extract text from resume file -> {file_path}")
            
            # Extract information
            try:
                result = self.llm.invoke(resume_text["text"])
                # rsp = self.flatten_meta(result)
                resume_text["metadata"] = result
                outputs.append(resume_text)
                # return result
            except Exception as e:
                logger.warning(f"Error during parsing: {e}")
                # return None
                resume_text["metadata"] = ""
                outputs.append(resume_text)

        return outputs
            
    def save_to_json(self, doc_metadata: TaxInvoiceModel, output_path: str):
        """Save extracted metadata to JSON file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                if isinstance(doc_metadata, dict):
                    f.write(json.dumps(doc_metadata, indent=2))
                else:
                    f.write(doc_metadata.model_dump_json(indent=2))
            print(f"Resume metadata saved to {output_path}")
        except Exception as e:
            print(f"Error saving to JSON: {e}")
    
    def print_document_summary(self, doc_metadata: TaxInvoiceModel):
        """Print a formatted summary of the extracted resume data"""
        print("\n" + "="*50)
        # Personal Information
        print(f"\n👤 STATEMENT INFORMATION:")
        print(f"account_info: {doc_metadata.model_dump()}")
        
        print("\n" + "="*50)

    def flatten_meta(self, doc_metadata: TaxInvoiceModel):
        """Print a formatted summary of the extracted resume data"""
        # flatten_meta = {}        
        # Personal Information
        print(f"\n👤 STATEMENT INFORMATION:")
        flatten_meta = doc_metadata.model_dump()

        return flatten_meta

async def adocumentParser(resume_file):
    """Example usage of the bank statement Parser"""
    
    # Example 1: Using Ollama (make sure Ollama is running with the model installed)
    model_provider = os.getenv("LLM_PROVIDER", "ollama")
    model_name = os.getenv("MODEL_NAME", "gemma3")
    try:
        print("Initializing bank statement Parser with Ollama...")
        parser_ollama = TaxInvoiceParser(model_provider=model_provider, model_name=model_name)
            # Parse the resume
        # print(f"Parsing resume: {resume_file}")
        metadata = await parser_ollama.adocumentParser(resume_file)
        
    except Exception as e:
        print(f"Error: {e}")
        metadata = None

    return metadata

def documentParser(resume_file, name, doc_type):
    """Example usage of the bank statement Parser"""
    
    # Example 1: Using Ollama (make sure Ollama is running with the model installed)
    model_provider = os.getenv("LLM_PROVIDER", "ollama")
    model_name = os.getenv("MODEL_NAME", "gemma3")
    try:
        print(f"Initializing bank statement Parser with {model_provider} : {model_name}...")
        parser_ollama = TaxInvoiceParser(model_provider=model_provider, model_name=model_name)
            # Parse the resume
        # print(f"Parsing resume: {resume_file}")
        # metadata = parser_ollama.llm.invoke(query = "", image = resume_file) #FOR DEBUG
        #
        metadata = parser_ollama.documentParser(resume_file, name, doc_type)
        
    except Exception as e:
        print(f"Error: {e}")
        metadata = None

    return metadata

# Example usage
async def amain():
    """Example usage of the bank statement Parser"""
    
    # Example 1: Using Ollama (make sure Ollama is running with the model installed)
    print("Initializing bank statement Parser with Ollama...")
    # parser_ollama = TaxInvoiceParser(model_provider="ollama", model_name="gemma3")
    
    # Example 2: Using Gemini (make sure GOOGLE_API_KEY is set)
    parser_ollama = TaxInvoiceParser(model_provider="gemini", model_name = "gemini-1.5-flash")
    
    # Parse a resume file
    resume_file = "/home/gs/Downloads/Statement.pdf"  # Replace with your resume file path
    resume_file = "/home/gs/work/reconcile-ai/test_data/bank_statement.pdf"

    with open(resume_file, "rb") as f:
        resume_file = f.read()

    try:
        # Parse the resume
        print(f"Parsing resume: {resume_file}")
        metadata = await parser_ollama.adocumentParser(resume_file)
        
        # Display summary
        # parser_ollama.print_document_summary(metadata)
        # parser_ollama.flatten_meta(metadata)
        
        # Save to JSON
        output_file = "extracted_resume_data.json"
        parser_ollama.save_to_json(metadata, output_file)
        
    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    resume_file = "/home/gs/Downloads/reconcile_data/0525/New Microsoft Word Document.docx"  # Replace with your resume file path
    resume_file = "/home/gs/Downloads/reconcile_data/0525/New Microsoft Word Document1.pdf"
    resume_file = "/home/gs/Downloads/reconcile_data/0525/New Microsoft Word Document.docx"
    # resume_file = "/home/gs/Downloads/reconcile_data/0525/照片 28-05-25 16 39 22.png"

    import pickle, base64, io
    from PIL import Image
    # with open("images.pkl", "rb") as fo:
    #     resume_file2 = pickle.load(fo)
    #     resume_file2 = resume_file2[0]

    #     buf = BytesIO(resume_file2)
    #     pil_image = Image.open(buf)#.convert("RGB")

    #     resume_file = pil_image2base64(pil_image, base64_only=True)

    # pil_image = Image.open(resume_file)#.convert("RGB")

    # resume_file = pil_image2base64(pil_image, base64_only=True)

    # main()
    import asyncio
    # asyncio.run(adocumentParser(resume_file))

    with open(resume_file, 'rb') as fo:
        bt = fo.read()
    # a=documentParser(bt, resume_file, "pdf")
    # a=documentParser(bt, resume_file, "image")
    a=documentParser(bt, resume_file, "doc")
    # print()
    print(f"** a -> {a[0].keys()}")