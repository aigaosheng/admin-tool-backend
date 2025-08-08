import os
import sys
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()

import json
from typing import Dict, List, Optional, Union
from dataclasses import dataclass, asdict
import PyPDF2
import docx
from docling.document_converter import DocumentConverter, PdfFormatOption
from llms.llm_model import LlmModel
from audit_logger.audit_logger import logger_audit_handler as logger
from io import BytesIO
from docling.datamodel.base_models import DocumentStream
from pydantic import BaseModel

# Data Models
class AccountInfo(BaseModel):
    account_name: str
    account_no: str
    account_deposit: Optional[float]
    account_withdrawl: Optional[float]
    account_balance: float

class Transaction(BaseModel):
    date: Optional[str] = None
    description: Optional[str] = None
    withdrawl: Optional[float] = None
    deposit: Optional[float] = None
    balance: Optional[float] = None

class BankAccountModel(BaseModel):
    account_info: AccountInfo
    transactions: List[Transaction] = []

class BankStatementModel(BaseModel):
    accounts: List[BankAccountModel] = []

class BankStatementParser:
    def __init__(self, model_provider: str = "ollama", model_name: str = "gemma3"):
        """
        Initialize bank statement Parser with chosen model provider
        
        Args:
            model_provider: "ollama" or "gemini"
            model_name: Model name (for Ollama) or not used for Gemini
        """
        # system_prompt = "You are an expert of bank statement parser. Extract the following information from the resume text below and return it as a JSON object."
        system_prompt = "You are a robust bank statement parser that can extract structured data from any quality data source. If any transaction record can not been parsed, SKIP the record. The bank statement are given in the below. Return as a JSON object."

        self.model_provider = model_provider
        self.model_name = model_name
        self.llm = LlmModel(model_provider, model_name, system_prompt=system_prompt, response_data_model=BankStatementModel)

        logger.info(f"LLM model: {model_provider}, {model_name}")
        
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            if isinstance(file_path, bytes):
                buf = BytesIO(file_path)
                pdf_reader = PyPDF2.PdfReader(buf)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            else:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            logger.info(f"Success bank statement to text -> {file_path[:50]} -> {text[:50]}")
        except Exception as e:
            logger.warning(f"Error reading PDF: {e}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            if isinstance(file_path, bytes):
                buf = BytesIO(file_path)
                doc = docx.Document(buf)
            else:
                doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            logger.info(f"Success bank statement to text -> {file_path[:50]} -> {text[:50]}")
        except Exception as e:
            logger.warning(f"Error reading DOCX: {e}")
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            if isinstance(file_path, bytes):
                text = file_path.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            logger.info(f"Success bank statement to text -> {file_path[:50]} -> {text[:50]}")

            return text
        except Exception as e:
            logger.warning(f"Error reading TXT: {e}")
            return ""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from resume file based on extension"""
        converter = DocumentConverter()
        try:
            if isinstance(file_path, bytes):
                buf = BytesIO(file_path)
                source = DocumentStream(name="my_doc.pdf", stream=buf)
            else:
                source = file_path
            result = converter.convert(source)

            md = result.document.export_to_markdown()
            logger.info(f"Success bank statement to text -> {file_path}")
        except Exception as e:
            logger.warning(f"Error proess bank statement to text -> {file_path}")
            md = ""  
        return md
    
    async def adocumentParser(self, file_path: str) -> BankStatementModel:
        """
        Parse resume and extract metadata
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            BankStatementModel object with extracted information
        """
        # Extract text from resume
        # resume_text = self.extract_text(file_path)
        resume_text = self.extract_text_from_pdf(file_path)
        
        if not resume_text.strip():
            logger.warning(f"Could not extract text from resume file -> {file_path}")
        
        # Extract information
        try:
            result = await self.llm(resume_text)
            # rsp = self.flatten_meta(result)
            return result
        except Exception as e:
            logger.warning(f"Error during parsing: {e}")
            return BankStatementModel()
    
    def documentParser(self, file_path: str) -> BankStatementModel:
        """
        Parse resume and extract metadata
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            BankStatementModel object with extracted information
        """
        # Extract text from resume
        # resume_text = self.extract_text(file_path)
        resume_text = self.extract_text_from_pdf(file_path)
        
        if not resume_text.strip():
            logger.warning(f"Could not extract text from resume file -> {file_path}")
        
        # Extract information
        try:
            result = self.llm.invoke(resume_text)
            # rsp = self.flatten_meta(result)
            return result
        except Exception as e:
            logger.warning(f"Error during parsing: {e}")
            return BankStatementModel()
            
    def save_to_json(self, doc_metadata: BankStatementModel, output_path: str):
        """Save extracted metadata to JSON file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                if isinstance(doc_metadata, dict):
                    f.write(json.dumps(doc_metadata, indent=2))
                else:
                    f.write(doc_metadata.model_dump_json(indent=2))
            print(f"Resume metadata saved to {output_path}")
        except Exception as e:
            print(f"Error saving to JSON: {e}")
    
    def print_document_summary(self, doc_metadata: BankStatementModel):
        """Print a formatted summary of the extracted resume data"""
        print("\n" + "="*50)
        # Personal Information
        print(f"\n👤 STATEMENT INFORMATION:")
        print(f"account_info: {doc_metadata.model_dump()}")
        
        print("\n" + "="*50)

    def flatten_meta(self, doc_metadata: BankStatementModel):
        """Print a formatted summary of the extracted resume data"""
        # flatten_meta = {}        
        # Personal Information
        print(f"\n👤 STATEMENT INFORMATION:")
        flatten_meta = doc_metadata.model_dump()

        return flatten_meta

async def adocumentParser(resume_file):
    """Example usage of the bank statement Parser"""
    
    # Example 1: Using Ollama (make sure Ollama is running with the model installed)
    model_provider = os.getenv("LLM_PROVIDER", "ollama")
    model_name = os.getenv("MODEL_NAME", "gemma3")
    try:
        print("Initializing bank statement Parser with Ollama...")
        parser_ollama = BankStatementParser(model_provider=model_provider, model_name=model_name)
            # Parse the resume
        # print(f"Parsing resume: {resume_file}")
        metadata = await parser_ollama.adocumentParser(resume_file)
        
    except Exception as e:
        print(f"Error: {e}")
        metadata = {}

    return metadata

def documentParser(resume_file):
    """Example usage of the bank statement Parser"""
    
    # Example 1: Using Ollama (make sure Ollama is running with the model installed)
    model_provider = os.getenv("LLM_PROVIDER", "ollama")
    model_name = os.getenv("MODEL_NAME", "gemma3")
    try:
        print(f"Initializing bank statement Parser with {model_provider} : {model_name}...")
        parser_ollama = BankStatementParser(model_provider=model_provider, model_name=model_name)
            # Parse the resume
        # print(f"Parsing resume: {resume_file}")
        metadata = parser_ollama.documentParser(resume_file)
        
    except Exception as e:
        print(f"Error: {e}")
        metadata = {}

    return metadata

# Example usage
async def amain():
    """Example usage of the bank statement Parser"""
    
    # Example 1: Using Ollama (make sure Ollama is running with the model installed)
    print("Initializing bank statement Parser with Ollama...")
    # parser_ollama = BankStatementParser(model_provider="ollama", model_name="gemma3")
    
    # Example 2: Using Gemini (make sure GOOGLE_API_KEY is set)
    parser_ollama = BankStatementParser(model_provider="gemini", model_name = "gemini-1.5-flash")
    
    # Parse a resume file
    resume_file = "/home/gs/Downloads/Statement.pdf"  # Replace with your resume file path
    resume_file = "/home/gs/work/reconcile-ai/test_data/bank_statement.pdf"

    with open(resume_file, "rb") as f:
        resume_file = f.read()

    try:
        # Parse the resume
        print(f"Parsing resume: {resume_file}")
        metadata = await parser_ollama.adocumentParser(resume_file)
        
        # Display summary
        # parser_ollama.print_document_summary(metadata)
        # parser_ollama.flatten_meta(metadata)
        
        # Save to JSON
        output_file = "extracted_resume_data.json"
        parser_ollama.save_to_json(metadata, output_file)
        
    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    resume_file = "/home/gs/Downloads/Statement.pdf"  # Replace with your resume file path
    resume_file = "/home/gs/work/reconcile-ai/test_data/bank_statement.pdf"

    # main()
    import asyncio
    asyncio.run(adocumentParser(resume_file))
    # documentParser(resume_file)